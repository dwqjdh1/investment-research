# 先安装 docx 库
import subprocess
subprocess.check_call(['pip', 'install', 'python-docx', '-q'])

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import datetime

doc = Document()

# ── 页面设置 ──
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(12)
style.paragraph_format.line_spacing = 1.25

# ═══════════════════════════════════════════
# 封面
# ═══════════════════════════════════════════
for i in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('自然语言处理综合实践报告')
run.font.size = Pt(22)
run.font.bold = True
run.font.name = '黑体'

doc.add_paragraph()

info_items = [
    ('项目名称', 'AI智能投研助手'),
    ('学    院', '_______________'),
    ('专    业', '_______________'),
    ('姓    名', '_______________'),
    ('学    号', '_______________'),
    ('指导教师', '_______________'),
]
for label, value in info_items:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 2.0
    run = p.add_run(f'{label}：{value}')
    run.font.size = Pt(14)
    run.font.name = '宋体'

doc.add_page_break()

# ═══════════════════════════════════════════
# 一、项目介绍
# ═══════════════════════════════════════════
h1 = doc.add_heading('一、项目介绍', level=1)
for run in h1.runs:
    run.font.name = '黑体'
    run.font.size = Pt(14)

def add_body(text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(24)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(text)
    run.font.name = '宋体'
    run.font.size = Pt(12)
    return p

add_body('随着中国资本市场的快速发展，A股和港股上市公司数量众多，投资者在面对海量财务数据时往往难以快速获取关键信息并做出有效决策。传统的投资研究方式需要投资者逐一查阅财务报表、行业报告，耗时耗力且容易遗漏重要信息。')
add_body('本项目"AI智能投研助手"旨在利用大语言模型（LLM）的自然语言处理能力，结合实时金融数据接口，为用户提供一键式的深度投资研究报告生成服务。用户仅需输入股票代码或公司名称，系统即可自动获取最新的财务数据、行情走势和公司信息，通过AI模型进行分析整合，生成包含基本面分析、盈利能力分析、估值分析等内容的专业研究报告，并以可视化图表辅助展示。')

add_body('应用场景包括：')
scenes = [
    '个人投资者快速了解目标公司的基本面状况；',
    '投资顾问为客户提供初步的个股分析参考；',
    '金融学习者通过系统了解股票分析的核心指标和逻辑；',
    '量化研究人员快速获取结构化财务数据。',
]
for s in scenes:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(24)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(s)
    run.font.name = '宋体'
    run.font.size = Pt(12)

add_body('项目核心价值在于将AI自然语言处理技术与金融数据相结合，降低了投资研究的门槛，让普通投资者也能获得接近专业水平的分析报告。')

doc.add_page_break()

# ═══════════════════════════════════════════
# 二、项目成果展示
# ═══════════════════════════════════════════
h2 = doc.add_heading('二、项目成果展示', level=1)
for run in h2.runs:
    run.font.name = '黑体'
    run.font.size = Pt(14)

add_body('本项目已成功部署于阿里云ECS服务器，可通过公网直接访问使用。同时项目源代码托管于GitHub，方便后续维护和二次开发。')

add_body('GitHub仓库地址：https://github.com/dwqjdh1/investment-research')
add_body('在线访问地址：http://<ECS公网IP>:8501')

add_body('系统主要功能如下：')
features = [
    '智能搜索：支持A股6位代码和港股5位代码的自动识别，支持中文名称模糊搜索；',
    '热门股票一键分析：内置11只热门A股和港股，点击即可生成研报；',
    '深度研报生成：调用大语言模型，自动整合财务数据、估值指标和行情走势，生成结构化研究报告；',
    '四大可视化图表：营收利润趋势图、盈利能力趋势图（ROE/毛利率/净利率）、估值仪表盘（PE/PB）、历史价格走势图；',
    'API配置灵活性：支持OpenAI兼容接口，用户可自定义Base URL、Model和API Key，适配DeepSeek等各种大语言模型服务。',
]
for f in features:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(24)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run('• ' + f)
    run.font.name = '宋体'
    run.font.size = Pt(12)

doc.add_page_break()

# ═══════════════════════════════════════════
# 三、架构设计及系统实现
# ═══════════════════════════════════════════
h3 = doc.add_heading('三、架构设计及系统实现', level=1)
for run in h3.runs:
    run.font.name = '黑体'
    run.font.size = Pt(14)

# 3.1 系统架构
h3_1 = doc.add_heading('3.1 系统架构设计思路', level=2)
for run in h3_1.runs:
    run.font.name = '黑体'
    run.font.size = Pt(12)

add_body('系统采用经典的分层架构设计，自上而下分为展示层、业务逻辑层、数据服务层三个层次：')

layers = [
    '展示层（Presentation Layer）：基于Streamlit框架构建Web用户界面，负责用户交互、输入接收和结果展示。Streamlit提供丰富的组件生态，支持Plotly图表直接嵌入，并通过CSS自定义实现现代化主题。',
    '业务逻辑层（Business Logic Layer）：包含报告生成引擎（report_generator.py）和可视化引擎（visualizer.py）。报告生成引擎通过prompts.py中定义的专业提示词模板，将原始财务数据输入LLM进行分析和自然语言生成。可视化引擎使用Plotly创建交互式金融图表。',
    '数据服务层（Data Service Layer）：基于AKShare金融数据接口获取A股和港股实时行情、历史价格、财务报表数据。LLM客户端（llm_client.py）封装OpenAI兼容接口调用，配置管理模块（config.py）负责环境变量和API密钥管理。',
]
for l in layers:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(24)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run('（' + str(layers.index(l)+1) + '）' + l)
    run.font.name = '宋体'
    run.font.size = Pt(12)

add_body('数据流如下：用户在搜索框输入代码或名称 → 系统通过AKShare搜索匹配股票 → 并行获取行情、财务、历史价格三类数据 → 数据汇总后输入LLM生成分析报告 → 同时Plotly生成可视化图表 → 报告和图表一同返回前端展示。')

# 3.2 核心技术
h3_2 = doc.add_heading('3.2 核心技术原理及关键模块实现', level=2)
for run in h3_2.runs:
    run.font.name = '黑体'
    run.font.size = Pt(12)

add_body('（1）大语言模型调用（LLM Integration）')
add_body('本项目基于OpenAI兼容接口调用大语言模型，默认使用DeepSeek模型。在llm_client.py中封装了标准的Chat Completion调用和流式输出（chat_stream）两种模式。流式模式下设置stream=True，通过生成器逐块yield响应内容，Streamlit前端使用st.empty()占位符增量更新，实现报告的实时逐字展示，大幅缩短用户感知等待时间。提示词工程方面，通过prompts.py设计了专业的金融分析师角色提示词，引导模型按照特定的报告结构（公司概况、财务分析、SWOT分析、估值判断、投资建议）进行输出。')

add_body('（2）金融数据获取（AKShare Integration）')
add_body('AKShare是一个开源的Python金融数据接口库。系统通过AKShare获取三类数据：实时行情（stock_zh_a_spot / stock_hk_spot_em），获取股票的当前价格、涨跌幅、成交量等；财务报表（A股使用stock_financial_abstract，港股使用stock_financial_hk_analysis_indicator_em），获取多年营收、净利润、ROE、毛利率、净利率、资产负债率、每股经营现金流等10项核心财务指标；历史价格（stock_zh_a_hist_tx / stock_hk_hist），获取复权后的日线价格序列用于趋势图绘制。对于港股，额外通过stock_hk_company_profile_em获取行业分类和公司概况信息，并通过stock_hk_financial_indicator_em获取总股本用于市值计算。系统实现了带TTL（300秒）的全市场缓存和单股缓存机制，并支持启动时后台预热。数据获取采用ThreadPoolExecutor并行执行，行情、财务、价格三类数据同时请求，互不阻塞。')

add_body('（3）数据可视化（Plotly）')
add_body('可视化模块（visualizer.py）使用Plotly创建四种交互式图表：营收利润趋势图使用双Y轴设计，柱状图展示营收，折线图展示净利润；盈利能力趋势图展示ROE、毛利率、净利率三条指标曲线的变化趋势；估值仪表盘使用Gauge组件展示当前市盈率（PE），并标注市净率（PB）和最新股价；历史价格走势图展示前复权收盘价曲线和20日均线，帮助识别技术趋势。图表生成与LLM报告生成通过ThreadPoolExecutor并行执行，四张图表在后台线程中同时渲染，不阻塞AI报告流的展示。')

add_body('（4）智能市场识别与性能优化')
add_body('系统通过代码长度自动判断市场类型：6位数字为A股代码，5位数字为港股代码。对于A股，根据代码前缀自动加上交易所标识（sh/sz）。对于港股代码，自动补齐前导零至5位。搜索功能支持代码精确匹配和名称模糊匹配两种模式。性能方面，实现了三项关键优化：LLM流式输出（stream=True），用户实时看到AI撰写进度；图表与LLM并行生成（ThreadPoolExecutor，4线程），消除串行等待；Prompt瘦身（财务历史从3季缩减至1季），减少输入token降低首字延迟。')

# 3.3 开发环境
h3_3 = doc.add_heading('3.3 开发环境', level=2)
for run in h3_3.runs:
    run.font.name = '黑体'
    run.font.size = Pt(12)

add_body('项目开发与部署的技术栈如下：')

# 建表
table = doc.add_table(rows=7, cols=2)
table.style = 'Table Grid'

env_data = [
    ('开发语言', 'Python 3.11'),
    ('Web框架', 'Streamlit（用于构建数据驱动的AI应用界面）'),
    ('数据获取', 'AKShare 1.14+（开源金融数据接口）'),
    ('大模型调用', 'OpenAI SDK 1.x（兼容DeepSeek等模型）'),
    ('数据可视化', 'Plotly 5.x + Kaleido（静态图表导出）'),
    ('数据分析', 'Pandas 2.x'),
    ('部署方式', 'Docker + Docker Compose，阿里云ECS（Ubuntu 22.04）'),
]
for i, (k, v) in enumerate(env_data):
    cell_k = table.cell(i, 0)
    cell_v = table.cell(i, 1)
    cell_k.text = k
    cell_v.text = v
    for cell in [cell_k, cell_v]:
        for p in cell.paragraphs:
            p.paragraph_format.line_spacing = 1.25
            for run in p.runs:
                run.font.name = '宋体'
                run.font.size = Pt(12)

doc.add_page_break()

# ═══════════════════════════════════════════
# 四、实验结果及效果改进
# ═══════════════════════════════════════════
h4 = doc.add_heading('四、实验结果及效果改进', level=1)
for run in h4.runs:
    run.font.name = '黑体'
    run.font.size = Pt(14)

add_body('在项目开发与部署过程中，对系统进行了全面的功能测试和问题排查。')

h4_1 = doc.add_heading('4.1 功能测试', level=2)
for run in h4_1.runs:
    run.font.name = '黑体'
    run.font.size = Pt(12)

add_body('（1）A股分析测试：以贵州茅台（600519）为例进行测试。系统成功获取了8期财务历史数据，所有10项财务指标（营收、净利润、EPS、BPS、ROE、毛利率、净利率、资产负债率等）完整返回。LLM生成了包含公司概览、财务分析、SWOT分析、估值判断和投资建议五个章节的完整分析报告。四个可视化图表均正常渲染，561个交易日的价格数据正确显示。')

add_body('（2）港股分析测试：以腾讯控股（00700）为例进行测试。修复后系统成功获取了9年完整的财务历史数据（2017-2025），10项财务指标全部齐全（营收7517.66亿、净利润2248.42亿、ROE 21.13%、毛利率56.21%等）。总股本（91.26亿股）和总市值（3.97万亿港元）正确计算展示。股价在行情API不可用时通过PE×EPS自动反推（434.76港元）。')

h4_2 = doc.add_heading('4.2 问题发现与改进', level=2)
for run in h4_2.runs:
    run.font.name = '黑体'
    run.font.size = Pt(12)

add_body('问题1：估值仪表盘生成报错（TypeError: got multiple values for keyword argument \'margin\'）。')
add_body('原因分析：在visualizer.py的create_valuation_gauge函数中，COMMON_LAYOUT字典已包含margin键，调用fig.update_layout(**COMMON_LAYOUT, margin=dict(...))时导致重复传参。')
add_body('解决方案：将update_layout调用拆分为两步，先应用COMMON_LAYOUT，再单独设置margin覆盖，确保了布局参数的正确应用。')

add_body('问题2：港股财务数据大量缺失——9年数据仅1行，10项指标缺6项。')
add_body('原因分析：原代码使用了stock_hk_financial_indicator_em接口，该接口仅返回1行TTM汇总数据，且列名与_HK_FIN_MAP中的映射不匹配（如API返回"股东权益回报率(%)"而代码期望"净资产收益率"、"销售净利率(%)" vs "净利率"），导致roe、gross_margin、net_margin、debt_ratio等关键指标全部miss。另外，该接口缺少日期列，历史趋势无法构建。')
add_body('解决方案：切换至stock_financial_hk_analysis_indicator_em接口，该接口返回9年完整年报数据，具备明确的REPORT_DATE列和标准英文列名（OPERATE_INCOME、HOLDER_PROFIT、ROE_AVG、GROSS_PROFIT_RATIO等）。新增_HK_ANALYSIS_MAP映射表实现列名到内部键的标准化转换。')

add_body('问题3：港股总股本和总市值始终为空。')
add_body('原因分析：stock_hk_company_profile_em接口（用于获取行业、公司介绍等）的返回列中不包含总股本字段，原代码在profile中查找total_shares永远返回None，导致市值计算被跳过。')
add_body('解决方案：在get_stock_info港股分支中增加stock_hk_financial_indicator_em调用，从已发行股本(股)字段提取总股本。同时增加股价容错：当行情API不可用时，使用PE×EPS反推最新股价，确保市值可正常计算。')

add_body('问题4：报告生成等待时间长，用户体验差。')
add_body('原因分析：原报告生成流程为串行三步——先获取数据（3-10秒），再画4张图表（1-2秒），最后等LLM输出完整2048 token报告（15-60秒），用户需等待总时长之和。LLM调用使用非流式模式，用户必须等全部内容生成完毕才能看到结果。')
add_body('解决方案：实施三项性能优化：(1) LLM改为流式输出（stream=True），使用st.empty()占位符增量渲染，用户实时看到文字逐步出现；(2) 图表与LLM并行执行（ThreadPoolExecutor，max_workers=4），图表渲染时间被LLM生成时间完全覆盖；(3) Prompt精简，财务历史从3季缩减至1季，减少输入token降低首字延迟。优化后用户感知等待时间显著缩短。')

h4_3 = doc.add_heading('4.3 部署测试', level=2)
for run in h4_3.runs:
    run.font.name = '黑体'
    run.font.size = Pt(12)

add_body('项目通过Docker容器化后部署到阿里云ECS（2 vCPU / 2GB内存 / Ubuntu 22.04），使用docker-compose进行编排管理。部署过程中遇到以下问题：')
deploy_issues = [
    'Git HTTP/2协议兼容性问题：使用git pull时出现"Error in the HTTP2 framing layer"错误，通过git config --global http.version HTTP/1.1切换至HTTP/1.1协议解决；',
    '安全组端口配置：初次部署后外部无法访问，经排查发现ECS默认未开放8501端口，在安全组入方向添加TCP 8501规则后问题解决；',
    'Docker Compose版本兼容：新版Docker使用docker compose（空格）而非docker-compose（横线）命令格式。',
]
for d in deploy_issues:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(24)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run('• ' + d)
    run.font.name = '宋体'
    run.font.size = Pt(12)

# ═══════════════════════════════════════════
# 五、实验总结
# ═══════════════════════════════════════════
h5 = doc.add_heading('五、实验总结', level=1)
for run in h5.runs:
    run.font.name = '黑体'
    run.font.size = Pt(14)

add_body('本项目成功实现了一个基于大语言模型和金融数据接口的AI智能投研助手。通过将Streamlit前端框架、AKShare数据接口、Plotly可视化引擎与大语言模型相结合，构建了一个功能完整、用户友好的投资研究工具。')

add_body('在技术层面，项目取得了以下成果：')

summaries = [
    '实现了A股和港股双市场的自动识别与数据获取，覆盖行情、财务、价格三类核心数据；',
    '设计了专业的金融分析师提示词，引导LLM生成结构化的深度研究报告；',
    '开发了四类交互式可视化图表，直观展示公司的营收趋势、盈利能力、估值水平和价格走势；',
    '解决了估值图表重复参数、港股数据不完整、日期列名兼容性等实际问题，增强了系统的鲁棒性；',
    '完成了Docker容器化部署，实现了从本地开发到云端上线的完整流程。',
]
for s in summaries:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(24)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run('（' + str(summaries.index(s)+1) + '）' + s)
    run.font.name = '宋体'
    run.font.size = Pt(12)

add_body('存在的不足与改进方向：目前港股财务数据的完整度仍受限于AKShare的接口覆盖范围，部分小型港股公司的财务指标可能缺失；报告的质量依赖于底层LLM的分析能力，不同模型之间的输出差异较大；系统尚不支持美股市场，未来可扩展三大市场全覆盖。此外，可考虑引入RAG技术，结合实时的行业研报和新闻资讯，进一步提升分析报告的深度和时效性。')

# ═══════════════════════════════════════════
# 六、参考文献
# ═══════════════════════════════════════════
h6 = doc.add_heading('参考文献', level=1)
for run in h6.runs:
    run.font.name = '黑体'
    run.font.size = Pt(14)

refs = [
    '[1] AKShare Documentation. https://akshare.akfamily.xyz/ [EB/OL].',
    '[2] Streamlit Documentation. https://docs.streamlit.io/ [EB/OL].',
    '[3] Plotly Python Documentation. https://plotly.com/python/ [EB/OL].',
    '[4] OpenAI API Documentation. https://platform.openai.com/docs/ [EB/OL].',
    '[5] DeepSeek API Documentation. https://platform.deepseek.com/docs [EB/OL].',
    '[6] Docker Documentation. https://docs.docker.com/ [EB/OL].',
    '[7] Abdin, M., et al. Phi-4 Technical Report. arXiv:2412.08905, 2024.',
    '[8] 阿里云ECS产品文档. https://help.aliyun.com/product/25365.html [EB/OL].',
]
for ref in refs:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(ref)
    run.font.name = '宋体'
    run.font.size = Pt(10.5)

# ── 保存 ──
output_path = r'C:\Users\dwqjdh\Desktop\新建文件夹 (3)\AI智能投研助手_实践报告.docx'
doc.save(output_path)
print(f'报告已保存到: {output_path}')
